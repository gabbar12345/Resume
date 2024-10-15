import tempfile
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openai
import ast
from home.constants import pdf_path,image_path
import os
from home.prompts import sy
from fpdf import FPDF
from home.prompts import *
from datetime import datetime
# from home.gcp_secrets import access_secret_version
import re



# openai.api_key =settings.API_KEY
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY')
def get_response(user_prompt,system_prompt=sy):
    response = openai.chat.completions.create(
    model="gpt-3.5-turbo",                    
    messages=[{"role": "system", "content": system_prompt,},
              {"role": "user", "content": user_prompt}],
    temperature=1,
    top_p=1
    )
    # print(response)
    result=response.choices[0].message.content
    return result

def safe_literal_eval(expression):
    try:
        # Attempt to evaluate the string as a Python literal
        result = ast.literal_eval(expression)
        return result
    except (ValueError, SyntaxError):
        # If evaluation fails, return the original string
        return expression

def formatedResponse(prompt):
    Body=get_response(user_prompt=prompt)
    changeBody = safe_literal_eval(Body)
    return changeBody

class PersonalDetails:
    def __init__(self, name, email, phone, linkedin):
        self.name = name
        self.email = email
        self.phone = phone
        self.linkedin = linkedin
        # self.photo_path = photo_path

class Chapter:
    def __init__(self, title, subtitle,year,styear, body):
        self.title = title
        self.subtitle = subtitle
        self.year=year
        self.styear=styear
        self.body = body

class AcademicDetail:
    def __init__(self, year, degree, institute_university, cgpa_percentage,styear):
        self.year = year
        self.styear = styear
        self.degree = degree
        self.institute_university = institute_university
        self.cgpa_percentage = cgpa_percentage

class Project:
    def __init__(self, title, description,year,styear):
        self.title = title
        self.description = description
        self.year=year
        self.styear=styear

class ResearchPaper:
    def __init__(self, title, authors, publication):
        self.title = title
        self.authors = authors
        self.publication = publication

class ResumePDF(FPDF):
    def __init__(self, personal_details, chapters, academic_details, professional_summary, skills, positions_of_responsibility, projects, research_papers,jobRole,achievement):
        super().__init__()
        self.personal_details = personal_details
        self.academic_details = academic_details
        self.chapters = chapters
        self.professional_summary = professional_summary
        self.skills = skills
        self.positions_of_responsibility = positions_of_responsibility
        self.projects = projects
        self.research_papers = research_papers
        self.jobRole = jobRole
        self.achievement=achievement
        self.header_added = False

    def header(self):
        if not self.header_added:
            self.set_fill_color(240, 240, 240)  # Lighter grey background
            self.rect(0, 0, 210, 48, 'F')  # Grey rectangle for header
            self.set_font('Arial', 'B', 24)
            self.set_text_color(0, 0, 100)  # Dark blue color for name
            self.cell(150, 15, self.personal_details.name.title(), 0, 1, 'L')
            self.set_font('Arial', 'I', 12)
            self.set_text_color(100, 100, 100)  # Grey color for subtext
            self.cell(150, 7, self.jobRole.title(), 0, 1, 'L')
            # self.image(self.personal_details.photo_path, 170, 5, 30, 30)
            self.ln(0)
            self.set_font('Arial', '', 10)
            self.set_text_color(0, 0, 0)  # Reset text color to black
            self.cell(150, 5, self.personal_details.email, 0, 1, 'L')
            self.cell(150, 5, self.personal_details.linkedin, 0, 1, 'L', link=self.personal_details.linkedin)
            self.cell(150, 5, self.personal_details.phone, 0, 1, 'L')
            self.ln(0)
            # Add red horizontal line at the bottom of the grey background
            self.set_draw_color(255, 0, 0)  # Set line color to red
            self.line(0, 49, 210, 49)  # Draw line at y=48 (bottom of grey background)
            self.set_draw_color(0, 0, 0)  # Reset draw color to black
            self.header_added = True

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

    def add_section_title(self, title):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, title, 'B', 1, 'L')
        self.ln(1)

    def add_chapter(self):
        self.add_section_title('EXPERIENCE')
        for chapter in self.chapters:
            self.set_font('Arial', 'B', 11)
            self.cell(130, 6, chapter.title, 0, 0, 'L')
            self.set_font('Arial', 'I', 10)
            if chapter.styear or chapter.year:
                self.cell(60, 6, chapter.subtitle + '('+ chapter.styear + " - " + chapter.year + ")", 0, 1, 'R')
            else:
                self.cell(60, 6, chapter.subtitle + chapter.styear + chapter.year, 0, 1, 'R')
            self.set_font('Arial', '', 10)
            self.multi_cell(0, 5, chapter.body['description'])
            for point in chapter.body['bullet_points']:
                self.cell(5)
                self.cell(0, 5, chr(149) + ' ' + point, 0, 1)
        self.ln(2)

    def add_academic_details(self):
        self.add_section_title('EDUCATION')
        for detail in self.academic_details:
            self.set_font('Arial', 'B', 10)
            self.cell(130, 6, detail.degree, 0, 0)
            self.set_font('Arial', 'I', 10)
            self.cell(60, 6, detail.styear+" - "+detail.year, 0, 1, 'R')
            self.set_font('Arial', '', 10)
            self.cell(130, 5, detail.institute_university, 0, 0)
            self.cell(60, 5, detail.cgpa_percentage + " CGPA", 0, 1, 'R')
        self.ln(2)

    def add_professional_summary(self):
        self.add_section_title('PROFESSIONAL SUMMARY')
        self.set_font('Arial', '', 10)
        self.multi_cell(0, 5, self.professional_summary)
        self.ln(1)

    def add_skills(self):
        # Add a section title
        self.add_section_title('SKILLS')
        # Set regular font for skills content
        self.set_font('Arial', '', 10) 
        # Join all skills with a separator and add them using MultiCell
        skills_text = ' | '.join(self.skills)
        # Width 0 means the cell width will be the remaining space to the right margin
        # Setting line height to 10 units
        self.multi_cell(0, 5, skills_text)
        
        # Add some vertical space after the skills section
        self.ln(2)

    def add_achievements(self):
        self.add_section_title('ACHIEVEMENT & CERTIFICATIONS')
        self.set_font('Arial', '', 10)
        for point in self.achievement:
            self.multi_cell(0, 5, chr(149) + ' ' + point, 0, 1)
            self.ln(2)

    def add_position_of_responsibility(self):
        self.add_section_title('POSITIONS OF RESPONSIBILITY')
        for position in self.positions_of_responsibility:
            self.set_font('Arial', 'B', 10)
            self.cell(130, 6, f"{position['position']} - {position['company']}", 0, 0)
            self.cell(60, 6, position['place'], 0, 1, 'R')
            self.set_font('Arial', '', 10)
            self.multi_cell(0, 5, position['impact'])
            self.ln(1)

    def add_projects(self):
        self.add_section_title('PROJECTS')
        for project in self.projects:
            self.set_font('Arial', 'B', 10)
            self.cell(130, 6, project.title, 0, 0)
            self.set_font('Arial', 'I', 10)
            self.cell(60, 6, project.styear + " - " + project.year, 0, 1, 'R')
            self.set_font('Arial', '', 10)
            self.multi_cell(0, 5, project.description)
        self.ln(2)

    def add_research_papers(self):
        self.add_section_title('RESEARCH PAPERS')
        for paper in self.research_papers:
            self.set_font('Arial', 'B', 10)
            self.cell(0, 6, paper.title, 0, 1)
            self.set_font('Arial', '', 10)
            self.cell(0, 5, f"Authors: {paper.authors}", 0, 1)
            self.cell(0, 5, f"Publication: {paper.publication}", 0, 1)
            self.ln(1)



def create_resume(fields):
    document = Document()

    # Title and Subtitle Formatting
    title_style = document.styles['Heading 1']
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(24)
    title_font.color.rgb = RGBColor(0, 0, 139)  # Dark Blue

    subtitle_style = document.styles['Heading 2']
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Arial'
    subtitle_font.size = Pt(14)
    subtitle_font.color.rgb = RGBColor(0, 128, 128)  # Teal

    body_style = document.styles['Normal']
    body_font = body_style.font
    body_font.name = 'Calibri'
    body_font.size = Pt(11)
    body_font.color.rgb = RGBColor(0, 0, 0)  # Black

    # Add Title
    document.add_heading(fields['name'], level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading(fields['contact information'], level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add Sections
    for section, content in fields['sections'].items():
        document.add_heading(section, level=2)
        for item in content:
            p = document.add_paragraph(item)
    return document



def save_resume(document, file_name=pdf_path):
    # Check if file exists and if so, that we can write to it
    if os.path.exists(file_name):
        if not os.access(file_name, os.W_OK):
            raise PermissionError(f"No write permission for file: {file_name}")

    try:
        # Save to a temporary file first
        temp_file = tempfile.NamedTemporaryFile(delete=False)
        temp_file_name = temp_file.name
        temp_file.close()

        # Saving document to a temporary file
        document.save(temp_file_name)

        # Replace the original file with the temporary file
        os.replace(temp_file_name, file_name)
        print(f"Document successfully saved to {file_name}")

    except PermissionError as e:
        print(f"PermissionError: {e}")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        # Clean up temporary file, if it exists
        if os.path.exists(temp_file_name):
            os.remove(temp_file_name)

def get_home_directory():
    return os.path.expanduser("~")

def generate_resume2(jobRole, request):

    personal_details = PersonalDetails(
        name=request.session.get('first_name') + " " + request.session.get('last_name'),
        email=request.session.get('email'),
        phone=request.session.get('phone'),
        linkedin=request.session.get('linkedin')
    )
    print(personal_details.name)
    today_date = datetime.now().strftime("%d %b %Y")

    academic_details = []
    for detail in request.session.get('education', []):
        endyear=(datetime.strptime(detail.get('college_end_date'),"%Y-%m-%d")).strftime("%d %b %Y")
        academic_details.append(AcademicDetail(
            year = "Present" if endyear == today_date else endyear,  
            styear=(datetime.strptime(detail.get('college_start_date'),"%Y-%m-%d")).strftime("%d %b %Y"),
            # (datetime.strptime(date_str, "%Y-%m-%d")).strftime("%b %Y").upper()
            # styear=detail.get('college_start_date'),
            degree=detail.get('degree'),
            institute_university=detail.get('college_name'),
            cgpa_percentage=detail.get('cgpa')
        ))
    
    chapters = []
    for chapter in request.session.get('experiences', []):
        employment_start_date = chapter.get('employment_start_date', '').strip()
        if employment_start_date:
            startyear=(datetime.strptime(employment_start_date,"%Y-%m-%d")).strftime("%d %b %Y")
        else:
            startyear=''
        employment_end_date = chapter.get('employment_end_date', '').strip()
        if employment_end_date:
            endyear=(datetime.strptime(employment_end_date,"%Y-%m-%d")).strftime("%d %b %Y")
        else:
            endyear=''
        job_det=chapter.get('job_details').split('\r')
        job_details = [jobdata.strip() for jobdata in job_det]
        chapters.append(Chapter(
            title=chapter.get('job_title'),
            subtitle=chapter.get('company_name'),
            year = "Present" if endyear == today_date else endyear,
            styear=startyear,
            # year=chapter.get('employment_end_date'),
            # styear=chapter.get('employment_start_date'),
            body={'description': chapter.get('job_description'),
                  'bullet_points': job_details if job_details else []}  # Split string into list
        ))  
    print(chapters)

    professional_summary=request.session.get('summary')
    skills=request.session.get('skills_languages').split(',')

    achievements=request.session.get('certification_1').split('\r')
    achievement = [achieve.strip() for achieve in achievements]
    
    projects = []
    for project in request.session.get('projects', []):
        endyear=(datetime.strptime(project.get('project_end_date'),"%Y-%m-%d")).strftime("%d %b %Y")
        projects.append(Project(
            title=project.get('project_title'),
            year = "Present" if endyear == today_date else endyear,
            styear=(datetime.strptime(project.get('project_start_date'),"%Y-%m-%d")).strftime("%d %b %Y"),
            description=project.get('project_description')
        ))


    research_papers = []
    for paper in request.session.get('research_papers', []):
        research_papers.append(ResearchPaper(
            title=paper.get('research_title'),
            authors=paper.get('research_authors'),
            publication=paper.get('research_publication')
        ))   
    # print(research_papers)

    positions_of_responsibility = []
    for position in request.session.get('positions_of_responsibility', []):
        positions_of_responsibility.append(position)
    
    # personal_details, academic_details, chapters, professional_summary, skills, positions_of_responsibility, projects, research_papers = get_sample_data()

    # Chapter Prompt Changes
    if chapters:
        for i in range(len(chapters)):
            chapter_prompt=resumePrompt.format(preData='chapters',data=chapters[i].body,job_role=jobRole)
            response=formatedResponse(chapter_prompt)
            chapters[i].body=response

    # Professional Summary Prompt Changes
    Prompt=resumePrompt.format(preData='professionalSummary',data=professional_summary,job_role=jobRole)
    summary_response=formatedResponse(Prompt)
    professional_summary=summary_response

    achiement_prompt=grammaticalPrompt.format(preData='Achievements & Certifications',data=achievement)
    achievement_response=formatedResponse(achiement_prompt)
    achievement=achievement_response

    for i in range(len(projects)):
        prompt=resumePrompt.format(preData='project description',data=projects[i].description,job_role=jobRole)
        project_response=formatedResponse(prompt)
        projects[i].description=project_response

   
    # Generate PDF
    pdf = ResumePDF(personal_details, chapters, academic_details, professional_summary, skills, positions_of_responsibility, projects, research_papers,jobRole,achievement)
    pdf.add_page()

    # Add light grey background for the entire page
    pdf.set_fill_color(250, 250, 250)  # Even lighter grey background
    pdf.rect(0, 48, 210, 297-40, 'F')  # 297 is A4 height, 40 is header height

    pdf.add_professional_summary()
    # for chapter in chapters:
    #     pdf.add_chapter(chapter)
    
    if chapters:
        pdf.add_chapter()
    pdf.add_projects()
    pdf.add_academic_details()
    pdf.add_skills()
    # pdf.add_position_of_responsibility()
    
    if research_papers:
        pdf.add_research_papers()
    if achievement[0]!='':
        pdf.add_achievements()
    


    return pdf
